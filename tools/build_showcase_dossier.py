"""Build the polished Synthia Showcase dossier from SHOWCASE_SUBMISSION.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "SHOWCASE_SUBMISSION.md"
OUTPUT = ROOT / "docs" / "showcase" / "Synthia_OpenAI_Showcase_Dossier.docx"
COVER = ROOT / "web" / "landing" / "assets" / "showcase-cover.png"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x20, 0x37, 0x48)
GRAY = RGBColor(0x55, 0x55, 0x55)
GOLD = RGBColor(0x7A, 0x5A, 0x00)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def mark_header_row(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    code = doc.styles.add_style("Showcase Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Consolas"
    code.font.size = Pt(9)
    code.font.color.rgb = DARK_BLUE
    code.paragraph_format.left_indent = Inches(0.22)
    code.paragraph_format.right_indent = Inches(0.22)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)

    callout = doc.styles.add_style("Showcase Callout", WD_STYLE_TYPE.PARAGRAPH)
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = INK
    callout.font.italic = True
    callout.paragraph_format.left_indent = Inches(0.25)
    callout.paragraph_format.right_indent = Inches(0.15)
    callout.paragraph_format.space_before = Pt(4)
    callout.paragraph_format.space_after = Pt(6)
    callout.paragraph_format.line_spacing = 1.208


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    for header in (section.header, section.even_page_header):
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.clear()

    for footer in (section.footer, section.even_page_footer):
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run("Page ")
        set_font(run, size=8, color=GRAY)
        field = OxmlElement("w:fldSimple")
        field.set(qn("w:instr"), "PAGE")
        p._p.append(field)


def clean_inline(text):
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = text.replace("**", "").replace("`", "").replace("*", "")
    return text


def style_list_paragraph(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.194)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.208


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("OPENAI SHOWCASE SUBMISSION DOSSIER")
    set_font(run, size=10.5, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("Synthia")
    set_font(run, size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Living Taxonomy for Scientific Memory")
    set_font(run, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("Prepared by Jean-Sebastien Beaulieu  |  SecuredMe  |  14 July 2026")
    set_font(run, size=9.5, color=GRAY, italic=True)

    if COVER.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(12)
        picture = p.add_run().add_picture(str(COVER), width=Inches(6.3))
        picture._inline.docPr.set(
            "descr",
            "Synthia Trace Lab showing the Aburria aburri taxonomy timeline and public evidence sources.",
        )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Traceable evidence. Explicit uncertainty. Human authority.")
    set_font(run, size=10.5, color=INK, bold=True)
    doc.add_page_break()


def add_markdown_body(doc, markdown):
    lines = markdown.splitlines()
    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        line = lines[index]
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Showcase Code")
                p.add_run("\n".join(code_lines))
                code_lines = []
            in_code = not in_code
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not line.strip():
            index += 1
            continue
        if line.startswith("# "):
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(clean_inline(line[3:]), level=1)
            index += 1
            continue
        if line.startswith("### "):
            doc.add_heading(clean_inline(line[4:]), level=2)
            index += 1
            continue
        if line.startswith("| "):
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                cells = [clean_inline(c.strip()) for c in lines[index].strip("|").split("|")]
                if not all(re.fullmatch(r"-+", cell.replace(":", "")) for cell in cells):
                    rows.append(cells)
                index += 1
            if rows:
                table = doc.add_table(rows=0, cols=len(rows[0]))
                table.style = "Table Grid"
                for ridx, row_values in enumerate(rows):
                    cells = table.add_row().cells
                    for cidx, value in enumerate(row_values):
                        cells[cidx].text = value
                        for run in cells[cidx].paragraphs[0].runs:
                            set_font(run, size=9.2, bold=(ridx == 0))
                        cells[cidx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                        cells[cidx].paragraphs[0].paragraph_format.space_after = Pt(0)
                        cells[cidx].paragraphs[0].paragraph_format.line_spacing = 1.15
                        if ridx == 0:
                            shade_cell(cells[cidx], "F4F6F9")
                columns = len(rows[0])
                widths = [9360 // columns] * columns
                widths[-1] += 9360 - sum(widths)
                if columns == 3:
                    widths = [1800, 3780, 3780]
                set_table_geometry(table, widths)
                mark_header_row(table.rows[0])
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        if re.match(r"^\d+\. ", line):
            parts = [re.sub(r"^\d+\. ", "", line).strip()]
            index += 1
            while index < len(lines) and lines[index].strip() and not re.match(
                r"^(#|>|\||- |\d+\. |```)", lines[index]
            ):
                parts.append(lines[index].strip())
                index += 1
            p = doc.add_paragraph(clean_inline(" ".join(parts)), style="List Number")
            style_list_paragraph(p)
            continue
        if line.startswith("- "):
            parts = [line[2:].strip()]
            index += 1
            while index < len(lines) and lines[index].strip() and not re.match(
                r"^(#|>|\||- |\d+\. |```)", lines[index]
            ):
                parts.append(lines[index].strip())
                index += 1
            p = doc.add_paragraph(clean_inline(" ".join(parts)), style="List Bullet")
            style_list_paragraph(p)
            continue
        if line.startswith(">"):
            parts = []
            while index < len(lines) and lines[index].startswith(">"):
                parts.append(clean_inline(lines[index][1:].strip()))
                index += 1
            p = doc.add_paragraph(style="Showcase Callout")
            p.add_run("\n".join(parts).strip())
            p_pr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F4F6F9")
            p_pr.append(shd)
            continue

        paragraph_lines = [line.strip()]
        index += 1
        while index < len(lines) and lines[index].strip() and not re.match(
            r"^(#|>|\||- |\d+\. |```)", lines[index]
        ):
            paragraph_lines.append(lines[index].strip())
            index += 1
        doc.add_paragraph(clean_inline(" ".join(paragraph_lines)))


def build():
    doc = Document()
    doc.settings.odd_and_even_pages_header_footer = True
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "Synthia OpenAI Showcase Submission Dossier"
    doc.core_properties.subject = "Living taxonomy and scientific memory"
    doc.core_properties.author = "Jean-Sebastien Beaulieu"
    doc.core_properties.keywords = "Synthia, Codex, taxonomy, scientific memory, showcase"
    add_cover(doc)
    add_markdown_body(doc, SOURCE.read_text(encoding="utf-8"))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
